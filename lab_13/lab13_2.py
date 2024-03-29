from tkinter import *
from tkinter import messagebox
from tkinter.ttk import Button, Combobox, Label
from lab13_pack.rectangle import Rectangle
from lab13_pack.triangle import Triangle
from lab13_pack.trapezoid import Trapezoid
import docx


class Window:
    def __init__(self, width, height, title="MyWindow", resizable=(False, False), icon=None):
        self.root = Tk()
        self.root.title(title)
        self.login_entry = Entry(self.root)
        self.age_entry = Entry(self.root)
        self.password_entry = Entry(self.root, show="*")

        if icon:
            self.root.iconbitmap(icon)
        self.shape = Combobox(self.root, values=('прямоугольник','треугольник','трапеция'), state="readonly")
        self.a_entry = Entry(self.root, width=10)
        self.b_entry = Entry(self.root, width=10)
        self.c_entry = Entry(self.root, width=10)
        self.h_entry = Entry(self.root, width=10)
        self.text1 = 'еее'
        self.doc = docx.Document()


    def run(self):
        self.draw_widgets()
        self.root.mainloop()

    def draw_widgets(self):
        #Текста-памятка
        Label(self.root, text='Для поиска площади введите:\n 1) ширину (a) и длинну (b) прямоугольника\n 2) основание (a)'
                              ' и высоту (h) треугольника\n 3) верхнее(a) и нижнее(b) основания и высоту(h)      \n трапеции',
              justify=LEFT, background='PaleTurquoise1').grid(row=1, column=0, sticky=W)
        Label(self.root,
              text='Для поиска радиуса описанной окружности введите:\n 1) ширину (a) и длинну (b) прямоугольника\n '
                 '2) длинны сторон(a, b, c) и высоту (h) треугольника\n 3) длинны сторон(a, b, c) и высоту(h) трапеции',
              justify=LEFT, background='PaleTurquoise2').grid(row=2, column=0, sticky=W)
        Label(self.root,
              text='Для поиска радиуса вписанной окружности введите:\n 1) ширину (a) и длинну (b) прямоугольника\n'
                     ' 2) длинны сторон(a, b, c) и высоту (h) треугольника\n 3) высоту(h) трапеции',
              justify=LEFT, background='PaleTurquoise3').grid(row=3, column=0, sticky=W)

        #Выбор фигуры из выпадающего списка
        Label(self.root, text="Выберите фигуру", justify=LEFT, background='CadetBlue1').grid(row=0, column=1, sticky=W)
        self.shape.grid(row=0, column=2, sticky=W)


        #Ввод переменных
        Label(self.root, text="Введите а", justify=LEFT, background='CadetBlue2').grid(row=1, column=1, sticky=W)
        self.a_entry.grid(row=1, column=2, sticky=W)

        Label(self.root, text="Введите b", justify=LEFT, background='CadetBlue2').grid(row=1, column=3, sticky=W)
        self.b_entry.grid(row=1, column=4, sticky=W)

        Label(self.root, text="Введите c", justify=LEFT, background='CadetBlue2').grid(row=2, column=1, sticky=W)
        self.c_entry.grid(row=2, column=2, sticky=W)

        Label(self.root, text="Введите h", justify=LEFT, background='CadetBlue2').grid(row=2, column=3, sticky=W)
        self.h_entry.grid(row=2, column=4, sticky=W)


        #Кнопки
        Button(self.root, text='Узнать площадь фигуры',  width=30, command=self.get_area).grid(row=1, column=5, sticky=W)

        Button(self.root, text='Узнать радиус описанной окружности фигуры', width=45, command=self.get_radius_c_circle).grid(row=2, column=5, sticky=W)

        Button(self.root, text='Узнать радиус описанной окружности фигуры', width=45, command=self.get_radius_i_circle).grid(row=3, column=5, sticky=W)
        Button(self.root, text='Запись', width=45,
               command=self.submit).grid(row=4, column=5, sticky=W)

    #Функции для кнопок
    # Площадь
    def get_area(self):
        s = self.shape.get()

        if s != 'прямоугольник' and s != 'треугольник' and s != 'трапеция':
            messagebox.showerror('Ошибка', 'Вы не выбрали фигуру')

        elif s == 'прямоугольник':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            res = str('Площадь прямоугольника со сторонами ' + str(a) +' и '+ str(b) + ' = ' + str(Rectangle.area(a, b)))

        elif s == 'треугольник':
            try:
                int(self.a_entry.get())
                int(self.h_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            h = int(self.h_entry.get())
            res = str('Площадь треугольника со стороной ' + str(a) + ' и высотой ' + str(h) + ' = ' + str(
                Triangle(1, a, 0, h, 0)))

        elif s == 'трапеция':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
                int(self.h_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            h = int(self.h_entry.get())
            res = str('Площадь трапеции со сторонами ' + str(a) + ', ' + str(b) + ' и высотой ' + str(h) + ' = ' + str(
                Trapezoid(1, a, b, 0, h)))
        messagebox.showinfo("Площадь фигуры", res)
        self.text1 = res


    #Радиус описанной окружности
    def get_radius_c_circle(self):
        s = self.shape.get()

        if s != 'прямоугольник' and s != 'треугольник' and s != 'трапеция':
            messagebox.showerror('Ошибка', 'Вы не выбрали фигуру')

        elif s == 'прямоугольник':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            res = str('Радиус описанной окружности прямоугольника со сторонами ' + str(a) +' и '+ str(b) + ' = ' +
                      str(Rectangle(2, a, b)))

        elif s == 'треугольник':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
                int(self.h_entry.get())
                int(self.c_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            c = int(self.c_entry.get())
            h = int(self.h_entry.get())
            res = str('Радиус описанной окружности треугольника со сторонами ' + str(a) + ', ' + str(b) + ', '+ str(c)
                      + ' и высотой ' + str(h) + ' = ' + str(Triangle( 2, a, b, c, h)))

        elif s == 'трапеция':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
                int(self.h_entry.get())
                int(self.c_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            c = int(self.c_entry.get())
            h = int(self.h_entry.get())
            res = str('Площадь трапеции со сторонами ' + str(a) + ', ' + str(b) + ', '+ str(c) + ' и высотой ' + str(h) + ' = ' + str(
                Trapezoid(2, a, b, c, h)))
        messagebox.showinfo("Радиус описанной окружности фигуры", res)
        self.text1 = res



    #Радиус вписанной окружности
    def get_radius_i_circle(self):
        s = self.shape.get()

        if s != 'прямоугольник' and s != 'треугольник' and s != 'трапеция':
            messagebox.showerror('Ошибка', 'Вы не выбрали фигуру')

        elif s == 'прямоугольник':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            res = str(
                'Радиус вписанной окружности прямоугольника со сторонами ' + str(a) + ' и ' + str(b) + ' = ' + str(
                    Rectangle(3, a, b)))

        elif s == 'треугольник':
            try:
                int(self.a_entry.get())
                int(self.b_entry.get())
                int(self.h_entry.get())
                int(self.c_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            a = int(self.a_entry.get())
            b = int(self.b_entry.get())
            c = int(self.c_entry.get())
            h = int(self.h_entry.get())
            res = str('Радиус вписанной окружности треугольника со сторонами ' + str(a) + ', ' + str(b) + ', '+ str(c)
                      + ' и высотой ' + str(h) + ' = ' + str(Triangle(3, a, b, c, h)))

        elif s == 'трапеция':
            try:
                int(self.h_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка', 'Вы ввели не корректное значение')
            h = int(self.h_entry.get())
            res = str('Площадь трапеции с высотой ' + str(h) + ' = ' + str(Trapezoid(3, 0, 0, 0, h)))
        messagebox.showinfo("Радиус вписанной окружности фигуры", res)
        self.text1 = res


    def submit(self):
        get1 = self.text1
        self.doc.add_paragraph(get1)
        self.doc.save('lab13_res.docx')
        text = []
        for paragraph in self.doc.paragraphs:
            text.append(paragraph.text)
            print('\n'.join(text))


if __name__ == "__main__":
    window = Window(500, 500, "Геометрические фигуры")
    window.run()
